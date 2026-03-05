import io
import os
import docx
import pdfplumber
import fitz  # PyMuPDF
from PIL import Image
from ocr import extract_text_from_pil_image

OCR_THRESHOLD = 50  # If extracted text has fewer chars than this, try OCR

def extract_images_from_pdf(file_bytes: bytes, output_dir: str) -> list[str]:
    """
    Extracts images from a PDF and saves them to output_dir.
    Returns a list of saved filenames.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    image_paths = []
    
    for page_index in range(len(doc)):
        page = doc[page_index]
        image_list = page.get_images(full=True)
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            ext = base_image["ext"]
            
            filename = f"pdf_image_p{page_index+1}_{img_index+1}.{ext}"
            filepath = os.path.join(output_dir, filename)
            
            with open(filepath, "wb") as f:
                f.write(image_bytes)
            image_paths.append(filename)
            
    return image_paths

def extract_images_from_docx(file_bytes: bytes, output_dir: str) -> list[str]:
    """
    Extracts images from a DOCX and saves them to output_dir.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    doc = docx.Document(io.BytesIO(file_bytes))
    image_paths = []
    img_count = 0
    
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_count += 1
            image_bytes = rel.target_part.blob
            # Guess extension from target_ref
            ext = rel.target_ref.split(".")[-1]
            filename = f"docx_image_{img_count}.{ext}"
            filepath = os.path.join(output_dir, filename)
            
            with open(filepath, "wb") as f:
                f.write(image_bytes)
            image_paths.append(filename)
            
    return image_paths

def read_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX file.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        return "\n\n".join(full_text)
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")

def read_pdf(file_bytes: bytes) -> tuple[str, str]:
    """
    Extracts text from a PDF file.
    Returns (text, method) where method is 'text' or 'ocr'.
    """
    text_content = []
    
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                
                if page_text and len(page_text.strip()) > 0:
                    text_content.append(page_text)
        
        extracted_text = "\n\n".join(text_content)
        
        if len(extracted_text.strip()) < OCR_THRESHOLD:
            print("Low text density in PDF, switching to OCR...")
            ocr_text_content = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    im = page.to_image(resolution=300).original
                    page_ocr_text = extract_text_from_pil_image(im)
                    ocr_text_content.append(page_ocr_text)
            
            return "\n\n".join(ocr_text_content), "ocr"

        return extracted_text, "text"

    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")

def process_document(filename: str, file_bytes: bytes, extract_images: bool = False, output_image_dir: str = None) -> dict:
    """
    Main entry point for document processing.
    Returns a dict with text, method, and optionally images.
    """
    filename_lower = filename.lower()
    result = {"text": "", "method": "", "images": []}
    
    if filename_lower.endswith(".docx"):
        result["text"] = read_docx(file_bytes)
        result["method"] = "docx_text"
        if extract_images and output_image_dir:
            result["images"] = extract_images_from_docx(file_bytes, output_image_dir)
            
    elif filename_lower.endswith(".pdf"):
        text, method = read_pdf(file_bytes)
        result["text"] = text
        result["method"] = method
        if extract_images and output_image_dir:
            result["images"] = extract_images_from_pdf(file_bytes, output_image_dir)
    else:
        raise ValueError("Unsupported file format. Please upload .pdf or .docx")
        
    return result
